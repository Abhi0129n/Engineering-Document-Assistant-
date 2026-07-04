import io
import os
from datetime import datetime

def generate_txt_bytes(messages, chat_title, document_name):
    """Generates a clean plain text export of the chat."""
    lines = []
    lines.append("ENGINEERING DOCUMENT ASSISTANT - CHAT EXPORT")
    lines.append("=" * 50)
    lines.append(f"Project:       Engineering Document Assistant")
    lines.append(f"Chat Session:  {chat_title or 'Current Chat'}")
    lines.append(f"Document Name: {document_name or 'No Active Document'}")
    lines.append(f"Export Date:   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 50)
    lines.append("")

    for msg in messages:
        role_label = "User" if msg["role"] == "user" else "Assistant"
        lines.append(f"{role_label}:")
        lines.append(msg["content"])
        if msg.get("role") == "assistant" and msg.get("confidence") is not None:
            lines.append(f"Confidence: {int(msg['confidence'])}%")
        lines.append("-" * 40)
        lines.append("")

    return "\n".join(lines).encode("utf-8")

def generate_md_bytes(messages, chat_title, document_name):
    """Generates a structured Markdown export preserving headings."""
    lines = []
    lines.append("# Engineering Document Assistant - Chat Export")
    lines.append("")
    lines.append(f"**Project:** Engineering Document Assistant  ")
    lines.append(f"**Chat Session:** {chat_title or 'Current Chat'}  ")
    lines.append(f"**Document Name:** {document_name or 'No Active Document'}  ")
    lines.append(f"**Export Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ")
    lines.append("")
    lines.append("---")
    lines.append("")

    for msg in messages:
        if msg["role"] == "user":
            lines.append("### User")
            lines.append(msg["content"])
        else:
            lines.append("### Assistant")
            lines.append(msg["content"])
            if msg.get("confidence") is not None:
                lines.append("")
                lines.append(f"*Confidence: {int(msg['confidence'])}%*")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines).encode("utf-8")

def generate_docx_bytes(messages, chat_title, document_name):
    """Generates a professional DOCX export using python-docx."""
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_header.add_run("L&T Construction | Confidential").font.size = Pt(8.5)

    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Engineering Document Assistant")
    run_title.font.name = "Helvetica"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(45, 125, 210) # Slate Blue
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("Chat Session Export Log")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(142, 142, 147)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Shading Accent 1'
    
    metadata = [
        ("Project Name", "Engineering Document Assistant"),
        ("Chat Title", chat_title or "Current Conversation"),
        ("Document Source", document_name or "No Active Document"),
        ("Generated On", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    ]
    
    for i, (key, value) in enumerate(metadata):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(key).font.bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph("") # Spacing
    doc.add_paragraph("---") # Divider

    for msg in messages:
        p_role = doc.add_paragraph()
        if msg["role"] == "user":
            run_role = p_role.add_run("User")
            run_role.font.bold = True
            run_role.font.size = Pt(11)
            run_role.font.color.rgb = RGBColor(45, 125, 210) # Blue
            
            p_content = doc.add_paragraph()
            run_content = p_content.add_run(msg["content"])
            run_content.font.size = Pt(10)
        else:
            run_role = p_role.add_run("Assistant")
            run_role.font.bold = True
            run_role.font.size = Pt(11)
            run_role.font.color.rgb = RGBColor(5, 150, 105) # Green
            
            p_content = doc.add_paragraph()
            run_content = p_content.add_run(msg["content"])
            run_content.font.size = Pt(10)
            
            if msg.get("confidence") is not None:
                p_conf = doc.add_paragraph()
                run_conf = p_conf.add_run(f"Confidence: {int(msg['confidence'])}%")
                run_conf.font.size = Pt(9)
                run_conf.font.italic = True
                run_conf.font.color.rgb = RGBColor(142, 142, 147)

        p_div = doc.add_paragraph()
        run_div = p_div.add_run("-" * 60)
        run_div.font.color.rgb = RGBColor(220, 220, 220)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_pdf_bytes(messages, chat_title, document_name):
    """Generates a professional PDF document of the chat history using ReportLab."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.pdfgen import canvas
    
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8.5)
            self.setFillColor(HexColor("#6B7280")) # Gray-500
            
            if self._pageNumber > 1:
                self.drawString(54, 782, "Engineering Document Assistant — Chat Session Export")
                self.setStrokeColor(HexColor("#E5E7EB"))
                self.setLineWidth(0.5)
                self.line(54, 775, 541, 775)
            
            self.setStrokeColor(HexColor("#E5E7EB"))
            self.setLineWidth(0.5)
            self.line(54, 48, 541, 48)
            
            self.drawString(54, 34, "Confidential — L&T Construction")
            page_text = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(541, 34, page_text)
            self.restoreState()

    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        leftMargin=54, 
        rightMargin=54, 
        topMargin=72, 
        bottomMargin=72
    )
    
    story = []
    
    styles = getSampleStyleSheet()
    
    normal = styles["Normal"]
    normal.fontName = "Helvetica"
    normal.fontSize = 10
    normal.leading = 14
    normal.textColor = HexColor("#1F2937") # Gray-800
    
    title_style = ParagraphStyle(
        "ChatTitleStyle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=HexColor("#2D7DD2"), # Accent Blue
        spaceAfter=12
    )
    
    user_style = ParagraphStyle(
        "ChatUserStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        textColor=HexColor("#2D7DD2") # Blue
    )
    
    assistant_style = ParagraphStyle(
        "ChatAssistantStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        textColor=HexColor("#059669") # Green
    )
    
    confidence_style = ParagraphStyle(
        "ChatConfidenceStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=HexColor("#6B7280"),
        spaceBefore=4
    )
    
    story.append(Paragraph("Engineering Document Assistant", title_style))
    story.append(Spacer(1, 10))
    
    meta_data = [
        [Paragraph("<b>Project Name:</b>", normal), Paragraph("Engineering Document Assistant", normal)],
        [Paragraph("<b>Chat Session:</b>", normal), Paragraph(str(chat_title or "Current Chat"), normal)],
        [Paragraph("<b>Document Name:</b>", normal), Paragraph(str(document_name or "No Active Document"), normal)],
        [Paragraph("<b>Generated Date:</b>", normal), Paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"), normal)]
    ]
    
    meta_table = Table(meta_data, colWidths=[120, 367])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor("#F9FAFB")),
        ('PADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LINEBELOW', (0, 0), (-1, -1), 0.5, HexColor("#F3F4F6")),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    
    story.append(meta_table)
    story.append(Spacer(1, 20))
    
    for idx, msg in enumerate(messages):
        msg_elements = []
        
        if msg["role"] == "user":
            heading = Paragraph(f"User:", user_style)
        else:
            heading = Paragraph(f"Assistant:", assistant_style)
        msg_elements.append(heading)
        msg_elements.append(Spacer(1, 4))
        
        content = Paragraph(msg["content"].replace("\n", "<br/>"), normal)
        msg_elements.append(content)
        
        if msg.get("role") == "assistant" and msg.get("confidence") is not None:
            conf = Paragraph(f"Confidence: {int(msg['confidence'])}%", confidence_style)
            msg_elements.append(conf)
            
        msg_elements.append(Spacer(1, 10))
        
        divider = Table([[""]], colWidths=[487])
        divider.setStyle(TableStyle([
            ('LINEBELOW', (0, 0), (-1, -1), 0.5, HexColor("#E5E7EB")),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
        ]))
        msg_elements.append(divider)
        msg_elements.append(Spacer(1, 12))
        
        story.append(KeepTogether(msg_elements))

    doc.build(story, canvasmaker=NumberedCanvas)
    return buffer.getvalue()

def generate_history_txt_bytes(all_chats, fetch_history_fn):
    """Compiles complete chat history database to plain text bytes."""
    lines = []
    lines.append("ENGINEERING DOCUMENT ASSISTANT - COMPLETE HISTORY EXPORT")
    lines.append("=" * 60)
    lines.append(f"Exported On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 60)
    lines.append("")

    for chat in all_chats:
        chat_id = chat["chat_id"]
        title = chat["title"]
        doc_name = chat["document_name"]
        messages = fetch_history_fn(chat_id)
        
        lines.append(f"Chat Session ID: {chat_id}")
        lines.append(f"Session Title:   {title}")
        lines.append(f"Document Source: {doc_name}")
        lines.append("-" * 50)
        
        for msg in messages:
            role_lbl = "User" if msg["role"] == "user" else "Assistant"
            lines.append(f"  [{msg.get('timestamp', '')}] {role_lbl}:")
            lines.append(f"  {msg['content']}")
            if msg.get("role") == "assistant" and msg.get("confidence") is not None:
                lines.append(f"  Confidence: {int(msg['confidence'])}%")
            lines.append("")
        
        lines.append("=" * 60)
        lines.append("")
        
    return "\n".join(lines).encode("utf-8")

def generate_history_md_bytes(all_chats, fetch_history_fn):
    """Compiles complete chat history database to Markdown bytes."""
    lines = []
    lines.append("# Engineering Document Assistant - Complete History Export")
    lines.append(f"*Exported On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
    lines.append("")
    lines.append("---")
    lines.append("")

    for chat in all_chats:
        chat_id = chat["chat_id"]
        title = chat["title"]
        doc_name = chat["document_name"]
        messages = fetch_history_fn(chat_id)
        
        lines.append(f"## Session: {title}")
        lines.append(f"**Document Name:** {doc_name}  ")
        lines.append(f"**Session ID:** `{chat_id}`  ")
        lines.append("")
        
        for msg in messages:
            role_lbl = "User" if msg["role"] == "user" else "Assistant"
            lines.append(f"### {role_lbl} *(at {msg.get('timestamp', '')})*")
            lines.append(msg["content"])
            if msg.get("role") == "assistant" and msg.get("confidence") is not None:
                lines.append("")
                lines.append(f"*Confidence: {int(msg['confidence'])}%*")
            lines.append("")
            
        lines.append("---")
        lines.append("")
        
    return "\n".join(lines).encode("utf-8")


def generate_history_docx_bytes(all_chats, fetch_history_fn):
    """Compiles complete chat history database to DOCX bytes."""
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_header.add_run("L&T Construction | Complete History Log").font.size = Pt(8.5)

    p_title = doc.add_paragraph()
    run_title = p_title.add_run("Engineering Document Assistant")
    run_title.font.name = "Helvetica"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(45, 125, 210)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run(f"Complete Database History Export - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(142, 142, 147)

    doc.add_paragraph("")
    doc.add_paragraph("---")

    for chat in all_chats:
        chat_id = chat["chat_id"]
        title = chat["title"]
        doc_name = chat["document_name"]
        messages = fetch_history_fn(chat_id)

        h2 = doc.add_paragraph()
        run_h2 = h2.add_run(f"Session: {title}")
        run_h2.font.bold = True
        run_h2.font.size = Pt(14)
        run_h2.font.color.rgb = RGBColor(45, 125, 210)

        p_meta = doc.add_paragraph()
        p_meta.add_run("Document Source: ").font.bold = True
        p_meta.add_run(f"{doc_name}\n")
        p_meta.add_run("Session ID: ").font.bold = True
        p_meta.add_run(f"{chat_id}")

        for msg in messages:
            p_role = doc.add_paragraph()
            role_lbl = "User" if msg["role"] == "user" else "Assistant"
            role_color = RGBColor(45, 125, 210) if msg["role"] == "user" else RGBColor(5, 150, 105)
            
            run_role = p_role.add_run(role_lbl)
            run_role.font.bold = True
            run_role.font.size = Pt(10)
            run_role.font.color.rgb = role_color
            
            p_content = doc.add_paragraph()
            p_content.add_run(msg["content"]).font.size = Pt(9.5)
            
            if msg.get("confidence") is not None:
                p_conf = doc.add_paragraph()
                run_conf = p_conf.add_run(f"Confidence: {int(msg['confidence'])}%")
                run_conf.font.size = Pt(8.5)
                run_conf.font.italic = True
                run_conf.font.color.rgb = RGBColor(142, 142, 147)

        doc.add_paragraph("---")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_history_pdf_bytes(all_chats, fetch_history_fn):
    """Compiles complete chat history database to PDF bytes."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether, PageBreak
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.pdfgen import canvas
    
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8.5)
            self.setFillColor(HexColor("#6B7280"))
            
            if self._pageNumber > 1:
                self.drawString(54, 782, "Engineering Document Assistant — Complete History Export")
                self.setStrokeColor(HexColor("#E5E7EB"))
                self.setLineWidth(0.5)
                self.line(54, 775, 541, 775)
            
            self.setStrokeColor(HexColor("#E5E7EB"))
            self.setLineWidth(0.5)
            self.line(54, 48, 541, 48)
            
            self.drawString(54, 34, "Confidential — L&T Construction")
            page_text = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(541, 34, page_text)
            self.restoreState()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        leftMargin=54, 
        rightMargin=54, 
        topMargin=72, 
        bottomMargin=72
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    normal = styles["Normal"]
    normal.fontName = "Helvetica"
    normal.fontSize = 9.5
    normal.leading = 13.5
    normal.textColor = HexColor("#1F2937")
    
    title_style = ParagraphStyle(
        "HistTitleStyle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=HexColor("#2D7DD2"),
        spaceAfter=12
    )
    
    session_title_style = ParagraphStyle(
        "HistSessionTitleStyle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=HexColor("#2D7DD2"),
        spaceBefore=12,
        spaceAfter=8
    )
    
    user_style = ParagraphStyle(
        "HistUserStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=14,
        textColor=HexColor("#2D7DD2")
    )
    
    assistant_style = ParagraphStyle(
        "HistAssistantStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=14,
        textColor=HexColor("#059669")
    )
    
    confidence_style = ParagraphStyle(
        "HistConfidenceStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=11,
        textColor=HexColor("#6B7280"),
        spaceBefore=4
    )
    
    story.append(Paragraph("Engineering Document Assistant", title_style))
    story.append(Paragraph(f"Complete Database History Export — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal))
    story.append(Spacer(1, 15))
    story.append(Table([[""]], colWidths=[487], style=[('LINEBELOW', (0,0), (-1,-1), 1.5, HexColor("#2D7DD2"))]))
    story.append(Spacer(1, 15))

    for idx, chat in enumerate(all_chats):
        chat_id = chat["chat_id"]
        title = chat["title"]
        doc_name = chat["document_name"]
        messages = fetch_history_fn(chat_id)
        
        if idx > 0:
            story.append(PageBreak())
            
        story.append(Paragraph(f"Session: {title}", session_title_style))
        
        meta_data = [
            [Paragraph("<b>Document Source:</b>", normal), Paragraph(str(doc_name), normal)],
            [Paragraph("<b>Session ID:</b>", normal), Paragraph(str(chat_id), normal)]
        ]
        meta_table = Table(meta_data, colWidths=[120, 367])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor("#F9FAFB")),
            ('PADDING', (0, 0), (-1, -1), 4),
            ('LINEBELOW', (0, 0), (-1, -1), 0.5, HexColor("#F3F4F6")),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 15))
        
        for msg in messages:
            msg_elements = []
            
            if msg["role"] == "user":
                heading = Paragraph(f"User:", user_style)
            else:
                heading = Paragraph(f"Assistant:", assistant_style)
            msg_elements.append(heading)
            msg_elements.append(Spacer(1, 4))
            
            content = Paragraph(msg["content"].replace("\n", "<br/>"), normal)
            msg_elements.append(content)
            
            if msg.get("role") == "assistant" and msg.get("confidence") is not None:
                conf = Paragraph(f"Confidence: {int(msg['confidence'])}%", confidence_style)
                msg_elements.append(conf)
                
            msg_elements.append(Spacer(1, 8))
            divider = Table([[""]], colWidths=[487], style=[('LINEBELOW', (0,0), (-1,-1), 0.5, HexColor("#F3F4F6"))])
            msg_elements.append(divider)
            msg_elements.append(Spacer(1, 8))
            
            story.append(KeepTogether(msg_elements))

    doc.build(story, canvasmaker=NumberedCanvas)
    return buffer.getvalue()
